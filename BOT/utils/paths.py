import pandas as pd
from selenium.common.exceptions import StaleElementReferenceException
import requests
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as Ec
import os
from docx import Document
from docx.shared import Inches
import io


class Path_Utils:

    def __init__(self,driver,timeout=10):
        self.driver = driver
        self.timeout = timeout

    def click_xpath(self,xpath):
        for _ in range(3):
            try:
                element = WebDriverWait(self.driver,timeout=self.timeout).until(
                    Ec.element_to_be_clickable((By.XPATH,xpath))
                )
                element.click()
                return True
            except StaleElementReferenceException as e:
                print(str(e))
                print("Stale element: Retrying...")
            except Exception as e:
                print(str(e))
                print(f"Exception during xpath of {xpath}")
        return False

    def send_keys_xpath(self,xpath,text:str):
        for _ in range(2):
            try:
                element = WebDriverWait(self.driver,timeout=self.timeout).until(
                    Ec.visibility_of_element_located((By.XPATH,xpath))
                )
                element.send_keys(text)
                return True
            except Exception as e:
                print(str(e))
                print(f"Exception during xpath {xpath}")
        return False

    def clear_xpath(self,xpath):
        for _ in range(2):
            try:
                element = WebDriverWait(self.driver,timeout=self.timeout).until(
                    Ec.element_to_be_clickable((By.XPATH,xpath))
                )
                element.clear()
                return True
            except Exception as e:
                print(str(e))
                print(f"Exception during xpath {xpath}")
        return False
            
    def get_element(self,xpath):
        for _ in range(2):
            try:
                element = WebDriverWait(self.driver,timeout=self.timeout).until(
                    Ec.visibility_of_element_located((By.XPATH,xpath))
                )
                return element
            except StaleElementReferenceException as e:
                print(str(e))
                print("Stale element Retrying.....")
            except Exception as e:
                print(str(e))

    def get_element_text(self,xpath):
            try:
                element = WebDriverWait(self.driver,timeout=self.timeout).until(
                    Ec.visibility_of_element_located((By.XPATH,xpath))
                )
                return str(element.text)
            except Exception as e:
                print(str(e))
                return "NO ELEMENT FOUND"
            
    def append_screenshot(self,filename,element):
            doc_path = f"{filename}.docx"

            img_url = element.get_attribute("src")
            response = requests.get(img_url)
            image_stream = io.BytesIO(response.content)

            if os.path.exists(doc_path):
                doc = Document(doc_path)
            else:
                doc = Document()

            doc.add_picture(image_stream, width=Inches(6))

            # Save the document
            doc.save(doc_path)

    def excel_write(df,path,sheetname):
        if not os.path.exists(path):
            with pd.ExcelWriter(path=path,engine='openpyxl') as writer:
                   df.to_excel(writer,sheet_name=sheetname)
        else:
            with pd.ExcelWriter(path=path,engine='openpyxl',mode='a') as writer:
                 df.to_excel(writer,sheet_name=sheetname)

# df = {
#      'Name':['Abhishek','Arun','Vijay'],
#      'Age':[21,23,24]
# }
# df = pd.DataFrame(df)
# Path_Utils.excel_write(df,path='Excel.xlsx',sheetname='OGa')


        

        

        
        