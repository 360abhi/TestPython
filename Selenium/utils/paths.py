from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as Ec
import os
from docx import Document
from docx.shared import Inches

class Path_Utils:

    def __init__(self,driver,timeout=10):
        self.driver = driver
        self.timeout = timeout

    def click_xpath(self,xpath):
        try:
            element = WebDriverWait(self.driver,timeout=self.timeout).until(
                Ec.element_to_be_clickable((By.XPATH,xpath))
            )
            element.click()
            return True
        except Exception as e:
            print(str(e))
            print(f"Exception during xpath of {xpath}")
            return False

    def send_keys_xpath(self,xpath,text:str):
        try:
            element = WebDriverWait(self.driver,timeout=self.timeout).until(
                Ec.visibility_of_element_located((By.XPATH,xpath))
            )
            element.send_keys(text)
            return True
        except Exception as e:
            print(str(e))
            print(f"Exception during xpath {xpath}")
            return False

    def clear_xpath(self,xpath):
        try:
            element = WebDriverWait(self.driver,timeout=self.timeout).until(
                Ec.element_to_be_clickable((By.XPATH,xpath))
            )
            element.clear()
            return True
        except Exception as e:
            print(str(e))
            print(f"Exception during xpath {xpath}")
            return False
        
    def get_element(self,xpath):
        try:
            element = WebDriverWait(self.driver,timeout=self.timeout).until(
                Ec.visibility_of_element_located((By.XPATH,xpath))
            )
            return element
        except Exception as e:
            print(str(e))

    def get_element_text(self,xpath):
        try:
            element = WebDriverWait(self.driver,timeout=self.timeout).until(
                Ec.visibility_of_element_located((By.XPATH,xpath))
            )
            return str(element.text)
        except Exception as e:
            print(str(e))
            return "NO ELEMENT FOUND"

    def save_screenshot_to_doc(self, filename):
        # Prepare file names
        screenshot_path = f"{filename}.png"
        doc_path = f"{filename}.docx"

        # Take screenshot
        self.driver.save_screenshot(screenshot_path)

        # Load or create Word document
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            doc = Document()

        # Add screenshot image to the document
        doc.add_paragraph(f"Screenshot added: {screenshot_path}")
        doc.add_picture(screenshot_path, width=Inches(6))  
        doc.add_paragraph("")  

        # Save the document
        doc.save(doc_path)
        os.remove(screenshot_path)

        print(f"Screenshot saved to {screenshot_path} and appended to {doc_path}")


    def append_screenshot(self,filename):
        screenshot_path = f"{filename}.png"
        doc_path = f"{filename}.docx"

        self.driver.save_screenshot(screenshot_path)

        if os.path.exists(filename):
            doc = Document(doc_path)
        else:
            doc = Document()

        doc.add_paragraph(f"Screenshot Added {screenshot_path}")
        doc.add_picture(screenshot_path,width=Inches(6))
        doc.add_paragraph("")


        doc.save(doc_path)
        os.remove(screenshot_path)
        print(f"Screenshot appended to {doc_path}")
